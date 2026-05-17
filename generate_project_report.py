from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "Rapport_Projet_Mini_Helpdesk.docx"
FALLBACK_OUTPUT_PATH = OUTPUT_DIR / "Rapport_Projet_Mini_Helpdesk_v2.docx"


PRIMARY = RGBColor(19, 66, 96)
ACCENT = RGBColor(20, 108, 96)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 110, 128)
LIGHT_FILL = "EAF3F6"
HEADER_FILL = "D6E7EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_run(run, *, bold=False, size=11, color=TEXT, font="Aptos") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def style_paragraph(paragraph, *, space_after=6, space_before=0, line_spacing=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = line_spacing


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=10)
    run = p.add_run("RAPPORT DE PROJET")
    style_run(run, bold=True, size=24, color=PRIMARY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=6)
    run = p.add_run("Mini Helpdesk Platform for Managing Support Tickets")
    style_run(run, bold=True, size=19, color=ACCENT)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=20)
    run = p.add_run("Solution cloud open source deployee en architecture PaaS")
    style_run(run, size=12, color=MUTED)

    info = [
        ("Module", "Cloud Computing"),
        ("Type de projet", "Application web de gestion de tickets"),
        ("Architecture finale", "Frontend Netlify + Backend Render + Render Postgres"),
        ("Mode de deploiement", "PaaS"),
        ("Membres", "Oussama EL KASMI\nMohammed LAMRINI\nAmin ATTAFI"),
        ("Date", date.today().strftime("%d/%m/%Y")),
    ]

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in info:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        for idx, cell in enumerate(row):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, space_after=0)
                for run in paragraph.runs:
                    style_run(run, bold=idx == 0, size=10, color=PRIMARY if idx == 0 else TEXT)
        set_cell_shading(row[0], HEADER_FILL)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=0)
    run = p.add_run(
        "Ce document presente le contexte, les choix techniques, l'architecture, "
        "les fonctionnalites, le deploiement et les limites du projet."
    )
    style_run(run, size=11, color=MUTED)

    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_paragraph()
    style_paragraph(heading, space_before=8 if level == 1 else 4, space_after=4)
    run = heading.add_run(text)
    style_run(
        run,
        bold=True,
        size=16 if level == 1 else 13,
        color=PRIMARY if level == 1 else ACCENT,
    )


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    style_paragraph(p, space_after=6, line_spacing=1.2)
    run = p.add_run(text)
    style_run(run, size=11, color=TEXT)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph()
        style_paragraph(p, space_after=3, line_spacing=1.1)
        run = p.add_run(f"- {item}")
        style_run(run, size=10.8, color=TEXT)


def add_numbered(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = document.add_paragraph()
        style_paragraph(p, space_after=3, line_spacing=1.1)
        run = p.add_run(f"{index}. {item}")
        style_run(run, size=10.8, color=TEXT)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, space_after=0)
            for run in paragraph.runs:
                style_run(run, bold=True, size=10, color=PRIMARY)

    for row_values in rows:
        row = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(row, row_values)):
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, space_after=0, line_spacing=1.1)
                for run in paragraph.runs:
                    style_run(run, bold=idx == 0, size=10, color=PRIMARY if idx == 0 else TEXT)

    document.add_paragraph()


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=0)
    run = p.add_run("Mini Helpdesk Platform - Rapport de projet")
    style_run(run, size=9, color=MUTED)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    add_footer(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)

    add_title_page(document)

    add_heading(document, "1. Resume du projet")
    add_body_paragraph(
        document,
        "Le projet consiste en la conception et la realisation d'une mini plateforme de helpdesk "
        "permettant de gerer des tickets de support. L'application offre une interface web pour "
        "creer des tickets, suivre leur traitement, attribuer des priorites, ajouter des commentaires "
        "et administrer les roles des utilisateurs. Elle est basee sur une pile open source composee "
        "de React, FastAPI, PostgreSQL et Docker.",
    )
    add_body_paragraph(
        document,
        "La version finale a ete deployee dans le cloud avec une architecture PaaS: le frontend est "
        "heberge sur Netlify, le backend sur Render et la base de donnees sur Render Postgres. "
        "Ce choix permet de repondre a la demande d'une solution cloud complete, sans gerer "
        "manuellement des machines virtuelles ou une infrastructure IaaS.",
    )

    add_heading(document, "2. Contexte et problematique")
    add_body_paragraph(
        document,
        "Dans de nombreux environnements, les demandes de support sont encore gerees de maniere "
        "informelle, par messages disperses ou sans suivi clair. Cela cree des pertes d'information, "
        "un manque de tracabilite et des difficultes pour prioriser le travail. Le projet propose "
        "donc une solution centralisee permettant de structurer la communication entre les utilisateurs "
        "et l'equipe de support.",
    )

    add_heading(document, "3. Objectifs du projet")
    add_bullets(
        document,
        [
            "Concevoir une application web de gestion de tickets simple, complete et fonctionnelle.",
            "Utiliser une pile open source moderne et cohérente.",
            "Mettre en place une architecture cloud basee sur des services PaaS.",
            "Assurer la separation des roles entre utilisateur, agent et administrateur.",
            "Proposer un deploiement heberge et une demonstration exploitable en ligne.",
        ],
    )

    add_heading(document, "4. Conformite avec la consigne")
    add_body_paragraph(
        document,
        "Le projet repond a la consigne de mise en place d'une solution open source cloud all in one. "
        "La pile technique utilise exclusivement des technologies open source cote application. "
        "L'architecture finale repose sur des plateformes cloud gerees, ce qui place clairement la "
        "solution dans le modele PaaS. Enfin, l'application constitue une solution complete comprenant "
        "frontend, backend, base de donnees, authentification, gestion des tickets, administration et "
        "deploiement automatise depuis GitHub.",
    )

    add_heading(document, "5. Choix technologiques")
    add_table(
        document,
        ["Composant", "Technologie", "Role"],
        [
            ["Frontend", "React + Vite", "Interface utilisateur web et navigation entre les pages"],
            ["Backend", "FastAPI", "Exposition de l'API REST et gestion de la logique metier"],
            ["Base de donnees", "PostgreSQL", "Stockage des utilisateurs, tickets, commentaires et pieces jointes"],
            ["Authentification", "JWT", "Gestion des sessions cote client et protection des routes"],
            ["Dev local", "Docker Compose", "Execution locale du frontend, du backend et de la base de donnees"],
            ["Versioning", "GitHub", "Gestion du code source et declenchement des redeploiements"],
        ],
    )

    add_heading(document, "6. Architecture de la solution")
    add_table(
        document,
        ["Service", "Plateforme", "Fonction"],
        [
            ["Frontend", "Netlify", "Hebergement de l'application React"],
            ["Backend API", "Render Web Service", "Execution de l'application FastAPI"],
            ["Base de donnees", "Render Postgres", "Hebergement de PostgreSQL"],
            ["Source", "GitHub", "Gestion du depot et deploiement automatique"],
        ],
    )
    add_numbered(
        document,
        [
            "L'utilisateur accede a l'interface web hebergee sur Netlify.",
            "Le frontend envoie ses requetes HTTP a l'API FastAPI hebergee sur Render.",
            "Le backend applique les regles de gestion et lit ou ecrit les donnees dans PostgreSQL.",
            "Les pieces jointes sont stockees dans le service backend pour le MVP, tandis que leurs metadonnees sont enregistrees en base.",
            "Chaque mise a jour du code sur GitHub peut declencher automatiquement un nouveau deploiement.",
        ],
    )

    add_heading(document, "7. Fonctionnalites realisees")
    add_bullets(
        document,
        [
            "Inscription et connexion des utilisateurs.",
            "Gestion des roles: user, agent et admin.",
            "Creation d'un ticket avec titre, description et priorite.",
            "Affichage de la liste des tickets avec filtres.",
            "Consultation detaillee d'un ticket.",
            "Mise a jour du statut et de la priorite d'un ticket.",
            "Attribution d'un ticket a un agent ou un administrateur.",
            "Ajout de commentaires sur un ticket.",
            "Televersement et suppression de pieces jointes.",
            "Tableau de bord administrateur avec statistiques simples et gestion des roles.",
        ],
    )

    add_heading(document, "8. Gestion des utilisateurs et securite")
    add_body_paragraph(
        document,
        "L'application utilise une authentification par jeton JWT. Lorsqu'un utilisateur se connecte, "
        "le backend genere un token qui est conserve cote frontend. Ce token est ensuite envoye dans "
        "les requetes protegees pour identifier l'utilisateur connecte. Les mots de passe sont haches "
        "avant d'etre stockes dans la base de donnees.",
    )
    add_body_paragraph(
        document,
        "Les autorisations dependent du role. Un utilisateur standard peut creer ses tickets, consulter "
        "ses propres demandes, commenter et ajouter des pieces jointes. Les agents et administrateurs "
        "peuvent gerer le workflow des tickets. L'administrateur dispose en plus d'un acces au tableau "
        "de bord et a la promotion des roles.",
    )

    add_heading(document, "9. Conception de la base de donnees")
    add_bullets(
        document,
        [
            "Table users: identite, email, mot de passe hache, role, date de creation.",
            "Table tickets: titre, description, statut, priorite, createur, assigne, dates de creation et de mise a jour.",
            "Table comments: contenu du commentaire, auteur, ticket associe et date.",
            "Table attachments: nom de fichier, URL de stockage, type MIME, taille, auteur du televersement et ticket associe.",
        ],
    )
    add_body_paragraph(
        document,
        "Ce schema garantit une structure simple mais suffisante pour gerer le coeur fonctionnel de "
        "la plateforme. Il permet aussi d'etendre facilement le projet avec de nouvelles fonctionnalites "
        "comme les notifications, l'historisation ou la gestion multi-equipes.",
    )

    add_heading(document, "10. Deploiement cloud et automatisation")
    add_body_paragraph(
        document,
        "Le frontend a ete deploye sur Netlify, tandis que le backend et la base de donnees ont ete "
        "deployes sur Render. Cette organisation permet une separation claire entre interface, logique "
        "metier et persistance des donnees. Elle simplifie egalement la mise en ligne du projet pour "
        "une demonstration en temps reel.",
    )
    add_body_paragraph(
        document,
        "Le projet beneficie aussi d'un workflow de deploiement automatise base sur GitHub. Lorsqu'une "
        "modification est poussee vers le depot, les plateformes cloud peuvent reconstruire et redeployer "
        "automatiquement l'application. Cela constitue une forme simple d'automatisation de livraison "
        "continue, adaptee au contexte du projet et a sa presentation.",
    )

    add_heading(document, "11. Limites actuelles et pistes d'amelioration")
    add_bullets(
        document,
        [
            "Le stockage des pieces jointes repose encore sur le systeme de fichiers du backend pour le MVP.",
            "Le projet ne gere pas encore les notifications par email.",
            "Le tableau de bord reste volontairement simple et pourrait etre enrichi avec des indicateurs plus avances.",
            "Un stockage objet dedie et plus persistant serait preferable pour une version de production.",
            "L'ajout de tests automatises et de controles CI renforcerait la qualite globale du pipeline.",
        ],
    )

    add_heading(document, "12. Conclusion")
    add_body_paragraph(
        document,
        "Ce projet a permis de realiser une application web complete de gestion de tickets, basee sur "
        "une pile open source et deployee dans le cloud en architecture PaaS. La plateforme couvre "
        "les besoins essentiels d'un mini helpdesk: authentification, gestion des roles, creation de tickets, "
        "suivi du workflow, commentaires, pieces jointes et administration.",
    )
    add_body_paragraph(
        document,
        "Au-dela de la partie developpement, le projet demontre aussi une bonne maitrise du deploiement "
        "cloud, de la communication entre frontend et backend, de la persistence des donnees et du lien "
        "entre GitHub et les plateformes d'hebergement. Il constitue ainsi une solution cloud coherente, "
        "fonctionnelle et presentable dans le cadre du module.",
    )

    add_heading(document, "13. Annexes")
    add_bullets(
        document,
        [
            "URL frontend: https://mini-helpdesk-frontend.netlify.app",
            "URL backend (health): https://mini-helpdesk-backend-bpm2.onrender.com/api/v1/health",
            "Compte administrateur de demonstration: cree par seed script dans la base distante.",
            "Deploiement local possible via Docker Compose pour disposer d'une solution de secours hors ligne.",
        ],
    )

    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    target_path = OUTPUT_PATH
    try:
        document.save(target_path)
    except PermissionError:
        target_path = FALLBACK_OUTPUT_PATH
        document.save(target_path)
    print(f"Saved report to: {target_path}")


if __name__ == "__main__":
    main()
